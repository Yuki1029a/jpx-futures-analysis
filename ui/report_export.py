"""Generate Word report as bytes for Streamlit download."""
from __future__ import annotations

import io
from datetime import date, timedelta
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from data.fetcher import get_volume_index, download_volume_excel, download_daily_oi_excel
from data.parser_volume import parse_volume_excel, parse_option_volume_excel, merge_volume_records
from data.parser_daily_oi import parse_daily_oi_excel, parse_daily_futures_oi_excel


# ---- helpers ----

def _set_cell(cell, text, bold=False, align="right", font_size=8, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "Meiryo"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    ea = rpr.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Meiryo"})
    rpr.append(ea)


def _shade(cell, hex_color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tcPr.append(shading)


def _fmt(n, sign=False):
    if n is None:
        return "-"
    return f"{n:+,.0f}" if sign else f"{n:,.0f}"


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Meiryo"
        rpr = run._element.get_or_add_rPr()
        ea = rpr.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Meiryo"})
        rpr.append(ea)


def _note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = "Meiryo"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _header_row(table, headers):
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, align="center", font_size=8)
        _shade(table.rows[0].cells[i], "2F5496")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


SESSION_LABEL = {
    "WholeDay": "立会(日中)", "Night": "立会(夜間)",
    "WholeDayJNet": "J-NET(日中)", "NightJNet": "J-NET(夜間)",
}
PRODUCT_NAME = {
    "NK225F": "日経225先物", "NK225MF": "日経225mini",
    "TOPIXF": "TOPIX先物", "NK225MicroF": "日経225micro",
}


# ---- public API ----

def build_report_bytes() -> tuple[bytes, str]:
    """Build Word report and return (docx_bytes, filename).

    Uses the latest available trade date from 202602 volume index.
    """
    entries = get_volume_index("202602")
    latest_entry = None
    latest_ds = None
    for e in reversed(entries):
        td = e.get("TradeDate", "").replace("/", "")
        if e.get("WholeDay"):
            latest_entry = e
            latest_ds = td
            break
    if not latest_entry:
        raise RuntimeError("No volume data found")

    report_date = date(int(latest_ds[:4]), int(latest_ds[4:6]), int(latest_ds[6:]))
    trade_label = f"{latest_ds[:4]}/{latest_ds[4:6]}/{latest_ds[6:]}"

    # ---- load data ----
    all_vol, session_data = [], {}
    for sk in ["WholeDay", "Night", "WholeDayJNet", "NightJNet"]:
        fpath = latest_entry.get(sk)
        if not fpath:
            continue
        content = download_volume_excel(fpath)
        rows = parse_volume_excel(content)
        all_vol.append(rows)
        session_data[sk] = rows

    merged = merge_volume_records(*all_vol) if all_vol else []

    prev_date = report_date - timedelta(days=1)
    while prev_date.weekday() >= 5:
        prev_date -= timedelta(days=1)

    oi_today = download_daily_oi_excel(report_date)
    futures_oi_today = parse_daily_futures_oi_excel(oi_today) if oi_today else []
    oi_records = parse_daily_oi_excel(oi_today) if oi_today else []

    opt_trades = []
    all_opt_raw = []
    for sk in ["WholeDay", "Night", "WholeDayJNet", "NightJNet"]:
        fpath = latest_entry.get(sk)
        if not fpath:
            continue
        content = download_volume_excel(fpath)
        rows = parse_option_volume_excel(content)
        for r in rows:
            opt_trades.append({"session": sk, "cm": r.contract_month,
                               "type": r.option_type, "strike": r.strike_price,
                               "pid": r.participant_id, "name": r.participant_name_en,
                               "vol": r.volume})
            all_opt_raw.append((sk, r))

    # ---- build doc ----
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.5))

    # Title
    title = doc.add_heading("先物・オプション手口分析", level=0)
    for run in title.runs:
        run.font.name = "Meiryo"
        rpr = run._element.get_or_add_rPr()
        rpr.append(rpr.makeelement(qn("w:rFonts"), {qn("w:eastAsia"): "Meiryo"}))
    p = doc.add_paragraph()
    run = p.add_run(f"取引日: {trade_label}  |  作成日: {date.today().isoformat()}")
    run.font.size = Pt(10)
    run.font.name = "Meiryo"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # --- S1: Futures volume ---
    _heading(doc, "1. 先物参加者別売買高")
    for product in ["NK225F", "TOPIXF"]:
        prod_label = PRODUCT_NAME.get(product, product)
        by_cm = defaultdict(list)
        for r in merged:
            if r.product == product:
                by_cm[r.contract_month].append(r)
        for cm in sorted(by_cm.keys()):
            cm_rows = sorted(by_cm[cm], key=lambda x: -x.volume)
            total = sum(r.volume for r in cm_rows)
            _heading(doc, f"{prod_label} ({cm}) - 合計: {_fmt(total)}枚", level=2)

            headers = ["参加者", "合計", "日中", "夜間", "J-NET", "J-NET%"]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            _header_row(t, headers)

            by_pid = defaultdict(lambda: {"name": "", "auction": 0, "jnet": 0, "day": 0, "night": 0})
            for sk, srows in session_data.items():
                is_jnet = "JNet" in sk
                for r in srows:
                    if r.product != product or r.contract_month != cm:
                        continue
                    by_pid[r.participant_id]["name"] = r.participant_name_en
                    if is_jnet:
                        by_pid[r.participant_id]["jnet"] += r.volume
                    else:
                        by_pid[r.participant_id]["auction"] += r.volume
            for r in cm_rows:
                if r.participant_id in by_pid:
                    by_pid[r.participant_id]["day"] = r.volume_day
                    by_pid[r.participant_id]["night"] = r.volume_night

            ranked = sorted(by_pid.items(), key=lambda x: -(x[1]["auction"] + x[1]["jnet"]))
            for pid, info in ranked[:20]:
                tot_p = info["auction"] + info["jnet"]
                if tot_p < 50:
                    continue
                jpct = info["jnet"] / tot_p * 100 if tot_p > 0 else 0
                row = t.add_row()
                _set_cell(row.cells[0], info["name"], align="left", font_size=7)
                _set_cell(row.cells[1], _fmt(tot_p))
                _set_cell(row.cells[2], _fmt(info["day"]))
                _set_cell(row.cells[3], _fmt(info["night"]))
                _set_cell(row.cells[4], _fmt(info["jnet"]))
                jc = RGBColor(0xC0, 0x39, 0x2B) if jpct > 50 else None
                _set_cell(row.cells[5], f"{jpct:.1f}%", color=jc)
            for row in t.rows:
                row.cells[0].width = Cm(5.5)
                for i in range(1, 6):
                    row.cells[i].width = Cm(2.2)
        doc.add_paragraph()

    # --- S2: Futures OI ---
    _heading(doc, "2. 先物建玉残高変化")
    _note(doc, f"{prev_date.strftime('%m/%d')} -> {report_date.strftime('%m/%d')}")
    h2 = ["商品", "限月", "出来高", "建玉残高", "前日比"]
    t2 = doc.add_table(rows=1, cols=len(h2))
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(t2, h2)
    for r in sorted(futures_oi_today, key=lambda x: (x.product, x.contract_month)):
        if r.contract_month not in ("2603", "2606"):
            continue
        row = t2.add_row()
        _set_cell(row.cells[0], PRODUCT_NAME.get(r.product, r.product), align="left")
        _set_cell(row.cells[1], r.contract_month, align="center")
        _set_cell(row.cells[2], _fmt(r.trading_volume))
        _set_cell(row.cells[3], _fmt(r.current_oi))
        cc = RGBColor(0xC0, 0x39, 0x2B) if r.net_change < 0 else (RGBColor(0x27, 0xAE, 0x60) if r.net_change > 0 else None)
        _set_cell(row.cells[4], _fmt(r.net_change, sign=True), color=cc)
    doc.add_paragraph()

    # --- S3: Option OI summary ---
    _heading(doc, "3. オプション建玉残高サマリー")
    h3 = ["限月", "CALL OI", "CALL変化", "PUT OI", "PUT変化", "P/C Ratio"]
    t3 = doc.add_table(rows=1, cols=len(h3))
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(t3, h3)
    for cm in ["2603", "2604", "2605", "2606"]:
        cm_recs = [r for r in oi_records if r.contract_month == cm]
        if not cm_recs:
            continue
        calls = [r for r in cm_recs if r.option_type == "CALL"]
        puts = [r for r in cm_recs if r.option_type == "PUT"]
        c_oi = sum(r.current_oi for r in calls)
        p_oi = sum(r.current_oi for r in puts)
        c_chg = sum(r.net_change for r in calls)
        p_chg = sum(r.net_change for r in puts)
        pcr = p_oi / c_oi if c_oi > 0 else 0
        row = t3.add_row()
        _set_cell(row.cells[0], cm, align="center", bold=True)
        _set_cell(row.cells[1], _fmt(c_oi))
        _set_cell(row.cells[2], _fmt(c_chg, sign=True),
                  color=RGBColor(0xC0, 0x39, 0x2B) if c_chg < 0 else (RGBColor(0x27, 0xAE, 0x60) if c_chg > 0 else None))
        _set_cell(row.cells[3], _fmt(p_oi))
        _set_cell(row.cells[4], _fmt(p_chg, sign=True),
                  color=RGBColor(0xC0, 0x39, 0x2B) if p_chg < 0 else (RGBColor(0x27, 0xAE, 0x60) if p_chg > 0 else None))
        _set_cell(row.cells[5], f"{pcr:.2f}",
                  color=RGBColor(0xC0, 0x39, 0x2B) if pcr > 1.2 else (RGBColor(0x27, 0xAE, 0x60) if pcr < 0.8 else None))
    doc.add_paragraph()

    # --- S4: Option OI movers ---
    _heading(doc, "4. オプション建玉変化 - 主要変動")
    for cm in ["2603", "2604", "2605"]:
        cm_recs = [r for r in oi_records if r.contract_month == cm]
        if not cm_recs:
            continue
        big = sorted(cm_recs, key=lambda x: abs(x.net_change), reverse=True)
        top = [r for r in big[:25] if abs(r.net_change) >= 10 or r.trading_volume >= 50]
        if not top:
            continue
        _heading(doc, f"限月: {cm}", level=2)
        h4 = ["種別", "行使価格", "建玉残高", "前日比", "出来高"]
        t4 = doc.add_table(rows=1, cols=len(h4))
        t4.style = "Table Grid"
        t4.alignment = WD_TABLE_ALIGNMENT.CENTER
        _header_row(t4, h4)
        for r in top:
            row = t4.add_row()
            tc = RGBColor(0x27, 0xAE, 0x60) if r.option_type == "CALL" else RGBColor(0xC0, 0x39, 0x2B)
            _set_cell(row.cells[0], r.option_type, align="center", color=tc)
            _set_cell(row.cells[1], _fmt(r.strike_price), align="center")
            _set_cell(row.cells[2], _fmt(r.current_oi))
            cc = RGBColor(0xC0, 0x39, 0x2B) if r.net_change < 0 else (RGBColor(0x27, 0xAE, 0x60) if r.net_change > 0 else None)
            _set_cell(row.cells[3], _fmt(r.net_change, sign=True), color=cc, bold=True)
            _set_cell(row.cells[4], _fmt(r.trading_volume))
        doc.add_paragraph()

    # --- S5: Large option trades ---
    _heading(doc, "5. 大口オプション取引 (50枚以上)")
    large = sorted([t for t in opt_trades if t["vol"] >= 50], key=lambda x: -x["vol"])
    h5 = ["セッション", "限月", "種別", "行使価格", "参加者", "売買高"]
    t5 = doc.add_table(rows=1, cols=len(h5))
    t5.style = "Table Grid"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(t5, h5)
    for tr in large[:40]:
        row = t5.add_row()
        sl = SESSION_LABEL.get(tr["session"], tr["session"])
        is_jnet = "JNet" in tr["session"]
        _set_cell(row.cells[0], sl, align="left", font_size=7,
                  color=RGBColor(0x8E, 0x44, 0xAD) if is_jnet else None)
        _set_cell(row.cells[1], tr["cm"], align="center")
        tc = RGBColor(0x27, 0xAE, 0x60) if tr["type"] == "CALL" else RGBColor(0xC0, 0x39, 0x2B)
        _set_cell(row.cells[2], tr["type"], align="center", color=tc)
        _set_cell(row.cells[3], _fmt(tr["strike"]), align="center")
        _set_cell(row.cells[4], tr["name"], align="left", font_size=7)
        _set_cell(row.cells[5], _fmt(tr["vol"]), bold=True)
    doc.add_paragraph()

    # --- S6: Set trades ---
    _heading(doc, "6. セット取引 / コンボ検出 (2603限月)")
    by_pid = defaultdict(list)
    for sk, r in all_opt_raw:
        if r.contract_month == "2603":
            by_pid[r.participant_id].append((sk, r))

    combo_found = False
    for pid, trades in sorted(by_pid.items()):
        calls = [(sk, t) for sk, t in trades if t.option_type == "CALL"]
        puts = [(sk, t) for sk, t in trades if t.option_type == "PUT"]
        tc_v = sum(t.volume for _, t in calls)
        tp_v = sum(t.volume for _, t in puts)
        if (tc_v < 30 and tp_v < 30) or not calls or not puts:
            continue
        name = trades[0][1].participant_name_en
        combo_found = True

        c_strikes = {t.strike_price for _, t in calls if t.volume >= 20}
        p_strikes = {t.strike_price for _, t in puts if t.volume >= 20}
        overlap = c_strikes & p_strikes
        strangles = [(cs, ps) for cs in c_strikes for ps in p_strikes if 0 < abs(cs - ps) <= 1000]
        structure = ""
        if overlap:
            structure = f"Straddle: {sorted(overlap)}"
        elif strangles:
            structure = "Strangle: " + ", ".join(f"C{cs}+P{ps}" for cs, ps in strangles)

        p = doc.add_paragraph()
        run = p.add_run(name)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = "Meiryo"
        run = p.add_run(f"  CALL={_fmt(tc_v)}  PUT={_fmt(tp_v)}")
        run.font.size = Pt(9)
        run.font.name = "Meiryo"
        if structure:
            run = p.add_run(f"  [{structure}]")
            run.font.size = Pt(8)
            run.font.name = "Meiryo"
            run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)

        detail = []
        for sk, t in sorted(calls, key=lambda x: -x[1].volume)[:5]:
            if t.volume >= 10:
                detail.append((sk, "CALL", t.strike_price, t.volume))
        for sk, t in sorted(puts, key=lambda x: -x[1].volume)[:5]:
            if t.volume >= 10:
                detail.append((sk, "PUT", t.strike_price, t.volume))
        if detail:
            dt = doc.add_table(rows=0, cols=4)
            dt.style = "Table Grid"
            for item in detail:
                row = dt.add_row()
                _set_cell(row.cells[0], SESSION_LABEL.get(item[0], item[0]), align="left", font_size=7)
                tc = RGBColor(0x27, 0xAE, 0x60) if item[1] == "CALL" else RGBColor(0xC0, 0x39, 0x2B)
                _set_cell(row.cells[1], item[1], align="center", font_size=7, color=tc)
                _set_cell(row.cells[2], _fmt(item[2]), align="center", font_size=7)
                _set_cell(row.cells[3], _fmt(item[3]), font_size=7, bold=True)

    if not combo_found:
        _note(doc, "該当なし")
    doc.add_paragraph()

    # --- S7: Matched C+P pairs ---
    _heading(doc, "7. C+P建玉対応ペア (Box/Conversion候補)")
    for cm in ["2603", "2604", "2605"]:
        cm_recs = [r for r in oi_records if r.contract_month == cm]
        by_strike = defaultdict(dict)
        for r in cm_recs:
            by_strike[r.strike_price][r.option_type] = r
        pairs = []
        for strike, types in sorted(by_strike.items()):
            if "CALL" in types and "PUT" in types:
                c, pv = types["CALL"], types["PUT"]
                if abs(c.net_change) >= 50 and abs(pv.net_change) >= 50:
                    mx = max(abs(c.net_change), abs(pv.net_change))
                    ratio = min(abs(c.net_change), abs(pv.net_change)) / mx if mx else 0
                    pairs.append((strike, c, pv, ratio))
        if not pairs:
            continue
        _heading(doc, f"限月: {cm}", level=2)
        h7 = ["行使価格", "CALL変化", "PUT変化", "CALL OI", "PUT OI", "一致率", "判定"]
        t7 = doc.add_table(rows=1, cols=len(h7))
        t7.style = "Table Grid"
        t7.alignment = WD_TABLE_ALIGNMENT.CENTER
        _header_row(t7, h7)
        for strike, c, pv, ratio in sorted(pairs, key=lambda x: -max(abs(x[1].net_change), abs(x[2].net_change))):
            row = t7.add_row()
            _set_cell(row.cells[0], _fmt(strike), align="center")
            _set_cell(row.cells[1], _fmt(c.net_change, sign=True))
            _set_cell(row.cells[2], _fmt(pv.net_change, sign=True))
            _set_cell(row.cells[3], _fmt(c.current_oi))
            _set_cell(row.cells[4], _fmt(pv.current_oi))
            _set_cell(row.cells[5], f"{ratio:.2f}", align="center")
            if ratio > 0.7:
                flag, fc = "MATCHED", RGBColor(0x8E, 0x44, 0xAD)
            elif ratio > 0.4:
                flag, fc = "Partial", RGBColor(0xF3, 0x9C, 0x12)
            else:
                flag, fc = "-", None
            _set_cell(row.cells[6], flag, align="center", color=fc, bold=True)
    doc.add_paragraph()

    _note(doc, "Data source: JPX (Japan Exchange Group)")
    _note(doc, "This report is generated automatically for analytical purposes only.")

    # serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"手口分析_{latest_ds}.docx"
    return buf.getvalue(), filename
